import logging

logger = logging.getLogger("resume_architect.doc_maker")

def generate_markdown(data: dict) -> str:
    name = data.get("Name", "Jane Doe")
    location = data.get("Location", "")
    phone = data.get("Phone", "")
    email = data.get("Email", "")
    linkedin = data.get("LinkedIn", "")
    summary = data.get("Summary", "")
    
    contact = " | ".join(c for c in [location, phone, email, linkedin] if c)
    
    md = f"# {name}\n\n{contact}\n\n"
    md += f"## Executive Summary\n\n{summary}\n\n"
    
    md += "## Experience & Projects\n\n"
    for exp in data.get("Experience", []):
        md += f"### {exp.get('Role', '')} | {exp.get('Company', '')} ({exp.get('Duration', '')})\n"
        for ach in exp.get("Achievements", []):
            md += f"- {ach}\n"
        md += "\n"
        
    md += "## Core Competencies\n\n"
    md += ", ".join(data.get("Skills", [])) + "\n\n"
    
    md += "## Education & Certifications\n\n"
    for edu in data.get("Education", []):
        md += f"- {edu}\n"
        
    return md

def generate_ats_text(data: dict) -> str:
    # Very clean standard text representation
    return generate_markdown(data)

def generate_docx_bytes(data: dict, theme: str) -> bytes:
    try:
        import docx
        from docx import Document
        from io import BytesIO
        
        doc = Document()
        
        # Add basic formatting
        h = doc.add_heading(data.get("Name", "Resume"), level=0)
        h.alignment = 1 # Center
        
        contact = " | ".join(c for c in [data.get("Location", ""), data.get("Phone", ""), data.get("Email", ""), data.get("LinkedIn", "")] if c)
        p = doc.add_paragraph(contact)
        p.alignment = 1
        
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(data.get("Summary", ""))
        
        doc.add_heading("Projects & Experience", level=1)
        for exp in data.get("Experience", []):
            p_exp = doc.add_paragraph()
            p_exp.add_run(f"{exp.get('Role', '')} ").bold = True
            p_exp.add_run(f"at {exp.get('Company', '')} ({exp.get('Duration', '')})").italic = True
            for ach in exp.get("Achievements", []):
                doc.add_paragraph(ach, style='List Bullet')
                
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(data.get("Skills", [])))
        
        doc.add_heading("Education & Certifications", level=1)
        for edu in data.get("Education", []):
            doc.add_paragraph(edu, style='List Bullet')
            
        f = BytesIO()
        doc.save(f)
        return f.getvalue()
    except Exception as e:
        logger.error(f"Failed to generate real docx, using fallback: {e}")
        # Return fallback text as bytes
        return generate_markdown(data).encode("utf-8")
